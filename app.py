import streamlit as st
import subprocess
import whisper
import tempfile
import os
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from datetime import timedelta


st.title("🎙️Transcripto")
st.subheader("Upload a video and get the transcription")

# Helper: Convert video to audio
def video_to_audio(video_path, audio_path="output.wav"):
    command = ["ffmpeg", "-y", "-i", video_path, audio_path]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return audio_path

# Cache Whisper model
@st.cache_resource
def get_audio_model(size):
    return whisper.load_model(size)

model_size = st.selectbox("Model size", ["tiny", "base", "small", "medium", "large"], index=1)

# Transcribe audio
def transcribe_audio(audio_path):
    model = get_audio_model(model_size)
    return model.transcribe(audio_path, fp16=False)

def format_timestamp(seconds):
    td = timedelta(seconds=seconds)
    total_seconds = td.total_seconds()
    hours, remainder = divmod(total_seconds, 3600)
    minutes, remainder = divmod(remainder, 60)
    secs, ms = divmod(remainder, 1)
    ms = int(ms * 1000)
    return f"{int(hours):02d}:{int(minutes):02d}:{int(secs):02d}.{ms:03d}"

# File generator for download
def download_file(transcription, base_filename, format_choice):
    if format_choice == "TXT":
        return transcription, f"{base_filename}_transcription.txt", "text/plain"

    elif format_choice == "DOCX":
        buffer = BytesIO()
        doc = Document()
        doc.add_heading("Transcription", 0)
        doc.add_paragraph(transcription)
        doc.save(buffer)
        buffer.seek(0)
        return buffer, f"{base_filename}_transcription.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif format_choice == "PDF":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        content = [Paragraph(line, styles["Normal"]) for line in transcription.split("\n")]
        doc.build(content)
        buffer.seek(0)
        return buffer, f"{base_filename}_transcription.pdf", "application/pdf"


# File uploader
uploaded_video = st.file_uploader("Upload a video", type=["mp4", "avi", "mov", "mkv"])

if uploaded_video:
    # Save uploaded video to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_file:
        tmp_file.write(uploaded_video.read())
        video_path = tmp_file.name

    # Convert video to audio
    audio_path = "audio.wav"
    video_to_audio(video_path, audio_path)

    # Run transcription
    with st.spinner("Transcribing..."):
        result = transcribe_audio(audio_path)

    # Display transcription
    transcription_lines = []
    for seg in result["segments"]:
        start = format_timestamp(seg["start"])
        end = format_timestamp(seg["end"])
        text = seg["text"].strip()
        transcription_lines.append(f"[{start} → {end}] {text}")
        st.markdown(f"**[{start} → {end}]** {text}")

    transcription = "\n".join(transcription_lines)

    # Format choice + download button
    format_choice = st.radio("Choose download format:", ["TXT", "DOCX", "PDF"])
    base_filename = os.path.splitext(uploaded_video.name)[0]

    file_data, file_name, mime = download_file(transcription, base_filename, format_choice)
    st.download_button(
        label=f"📥 Download {format_choice}",
        data=file_data,
        file_name=file_name,
        mime=mime
    )

    # Cleanup temp files
    os.remove(video_path)
    if os.path.exists(audio_path):
        os.remove(audio_path)